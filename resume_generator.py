from docx import Document
from docx.shared import Inches
from flask import Flask
from flask import url_for, render_template, request, redirect, flash, session
import requests
import json

app = Flask(__name__)

document = Document()

def make_a_docx():
    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.', style='ListBullet')
    paragraph.style = 'ListBullet'
    prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
    document.add_heading('The REAL meaning of the universe')
    document.add_heading('The role of dolphins', level=2)
    document.add_page_break()
    table = document.add_table(rows=2, cols=2)
    cell = table.cell(0, 1)
    cell.text = 'parrot, possibly dead'
    row = table.rows[1]
    row.cells[0].text = 'Foo bar to you.'
    row.cells[1].text = 'And a hearty foo bar to you too sir!'
    table.style = 'LightShading-Accent1'
    document.add_picture('picture.jpg', width=Inches(3.0))




def vk_api(method, **kwargs):
    api_request = 'https://api.vk.com/method/'+method + '?'
    api_request += '&'.join(['{}={}'.format(key, kwargs[key]) for key in kwargs])
    return json.loads(requests.get(api_request).text)

def collect_info():
    u_id = '1'
    flds = 'first_name,last_name,photo_50,city,career,education,'
    result = vk_api('users.get', user_ids=u_id, fields=flds, v='5.63')
    print(result)



@app.route('/')
def hello_world():
    return render_template('index.html')

@app.route('/access_error')
def access_error():
    pass


if __name__ == '__main__':
    #app.run(debug=True)
    collect_info()